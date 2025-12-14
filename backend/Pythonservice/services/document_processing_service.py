"""
Document Processing Service
Service để xử lý và trích xuất nội dung từ các loại tài liệu khác nhau
"""

import os
import json
from typing import Dict, Any, Tuple, Optional
import traceback

class DocumentProcessingService:
    """Service xử lý tài liệu và trích xuất text content"""

    def __init__(self):
        """Khởi tạo DocumentProcessingService"""
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx'],
            'doc': ['.doc'],
            'xlsx': ['.xlsx'],
            'xls': ['.xls'],
            'csv': ['.csv'],
            'txt': ['.txt']
        }

    def _detect_file_type(self, file_path: str, mime_type: str) -> str:
        """
        Detect file type từ MIME type hoặc file extension
        
        Args:
            file_path: Đường dẫn file
            mime_type: MIME type từ request
            
        Returns:
            String: detected file type (pdf, docx, xlsx, etc.)
        """
        # Mapping MIME types to file types
        mime_mapping = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'application/vnd.ms-excel': 'xls',
            'text/csv': 'csv',
            'text/plain': 'txt'
        }
        
        # Try MIME type first
        if mime_type in mime_mapping:
            return mime_mapping[mime_type]
        
        # Fallback to file extension
        _, ext = os.path.splitext(file_path)
        ext_lower = ext.lower()
        
        extension_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.xlsx': 'xlsx',
            '.xls': 'xls',
            '.csv': 'csv',
            '.txt': 'txt'
        }
        
        if ext_lower in extension_mapping:
            return extension_mapping[ext_lower]
        
        # Default to unknown
        return 'unknown'

    def extract_text_from_file(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Trích xuất text content từ file

        Args:
            file_path: Đường dẫn đến file
            file_type: Loại file (pdf, docx, xlsx, csv, etc.)

        Returns:
            Tuple[str, Dict]: (extracted_text, metadata)
        """
        try:
            if not os.path.exists(file_path):
                return "", {
                    "extraction_success": False,
                    "error": f"File không tồn tại: {file_path}",
                    "content_length": 0
                }

            file_type_lower = file_type.lower()

            # Detect file type từ MIME type hoặc extension
            detected_type = self._detect_file_type(file_path, file_type_lower)
            
            # Xử lý theo loại file đã detect
            if detected_type == 'pdf':
                return self._extract_pdf(file_path)
            elif detected_type in ['docx', 'doc']:
                return self._extract_docx(file_path)
            elif detected_type in ['xlsx', 'xls']:
                return self._extract_excel(file_path)
            elif detected_type == 'csv':
                return self._extract_csv(file_path)
            elif detected_type == 'txt':
                return self._extract_txt(file_path)
            else:
                return "", {
                    "extraction_success": False,
                    "error": f"Không hỗ trợ định dạng file: {file_type} (detected: {detected_type})",
                    "content_length": 0
                }

        except Exception as e:
            print(f"[Document Processing] Error extracting text from {file_path}: {str(e)}")
            traceback.print_exc()
            return "", {
                "extraction_success": False,
                "error": str(e),
                "content_length": 0
            }

    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Trích xuất text từ PDF file"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_content = ""

            for page in reader.pages:
                text_content += page.extract_text() + "\n"

            return text_content.strip(), {
                "extraction_success": True,
                "content_length": len(text_content),
                "pages": len(reader.pages),
                "file_size": os.path.getsize(file_path)
            }

        except ImportError:
            return "", {
                "extraction_success": False,
                "error": "PyPDF2 library not installed",
                "content_length": 0
            }
        except Exception as e:
            return "", {
                "extraction_success": False,
                "error": f"PDF extraction error: {str(e)}",
                "content_length": 0
            }

    def _extract_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Trích xuất text từ DOCX file"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = ""

            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content += " | ".join(row_text) + "\n"

            return text_content.strip(), {
                "extraction_success": True,
                "content_length": len(text_content),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_size": os.path.getsize(file_path)
            }

        except ImportError:
            return "", {
                "extraction_success": False,
                "error": "python-docx library not installed",
                "content_length": 0
            }
        except Exception as e:
            return "", {
                "extraction_success": False,
                "error": f"DOCX extraction error: {str(e)}",
                "content_length": 0
            }

    def _extract_excel(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Trích xuất text từ Excel file"""
        try:
            import pandas as pd

            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            text_content = ""
            sheets_info = []

            for sheet_name, df in excel_data.items():
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                text_content += df.to_string(index=False) + "\n"

                sheets_info.append({
                    "name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist()
                })

            return text_content.strip(), {
                "extraction_success": True,
                "content_length": len(text_content),
                "sheets": sheets_info,
                "total_sheets": len(sheets_info),
                "file_size": os.path.getsize(file_path)
            }

        except ImportError:
            return "", {
                "extraction_success": False,
                "error": "pandas library not installed",
                "content_length": 0
            }
        except Exception as e:
            return "", {
                "extraction_success": False,
                "error": f"Excel extraction error: {str(e)}",
                "content_length": 0
            }

    def _extract_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Trích xuất text từ CSV file"""
        try:
            import pandas as pd

            df = pd.read_csv(file_path)
            text_content = df.to_string(index=False)

            return text_content, {
                "extraction_success": True,
                "content_length": len(text_content),
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "file_size": os.path.getsize(file_path)
            }

        except ImportError:
            return "", {
                "extraction_success": False,
                "error": "pandas library not installed",
                "content_length": 0
            }
        except Exception as e:
            return "", {
                "extraction_success": False,
                "error": f"CSV extraction error: {str(e)}",
                "content_length": 0
            }

    def _extract_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Trích xuất text từ TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()

            return text_content, {
                "extraction_success": True,
                "content_length": len(text_content),
                "file_size": os.path.getsize(file_path)
            }

        except Exception as e:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text_content = f.read()

                return text_content, {
                    "extraction_success": True,
                    "content_length": len(text_content),
                    "file_size": os.path.getsize(file_path),
                    "encoding": "latin-1"
                }

            except Exception as e2:
                return "", {
                    "extraction_success": False,
                    "error": f"TXT extraction error: {str(e2)}",
                    "content_length": 0
                }


# Global instance
_document_processor = None

def get_document_processor() -> DocumentProcessingService:
    """Get global DocumentProcessingService instance"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessingService()
    return _document_processor